import os
import subprocess
import json
from typing import List, Dict, Any
import pypdf
import openpyxl
from docx import Document

from deepagents import create_deep_agent
from langchain_core.tools import tool
from langchain_openai import ChatOpenAI

# ----------------------------------------------------
# 1. Pipeline-Specific Custom Tools
# ----------------------------------------------------

@tool
def parse_raw_documents(file_configs: List[Dict[str, Any]], output_json_path: str = "workspace/master.json") -> str:
    """
    Parses complex, multi-format documentation (PDF page blocks, Excel sheets, Word files) 
    and normalizes register structures into a clean consolidated Master JSON file.
    """
    master_data = {}
    os.makedirs(os.path.dirname(output_json_path), exist_ok=True)
    
    for config in file_configs:
        path = config["path"]
        file_type = config["type"].lower()
        
        if not os.path.exists(path):
            return f"Error: Input file missing at location: {path}"
            
        master_data[path] = []
        
        if file_type == "pdf":
            pages = config.get("pages", [])
            with open(path, "rb") as f:
                reader = pypdf.PdfReader(f)
                start_page = pages[0] if len(pages) > 0 else 0
                end_page = pages[1] if len(pages) > 1 else len(reader.pages)
                
                for page_num in range(start_page, min(end_page + 1, len(reader.pages))):
                    text = reader.pages[page_num].extract_text()
                    master_data[path].append({"page": page_num, "content": text})
                    
        elif file_type in ["xlsx", "xlsm"]:
            sheets = config.get("sheets", [])
            wb = openpyxl.load_workbook(path, data_only=True)
            target_sheets = sheets if sheets else wb.sheetnames
            
            for sheet_name in target_sheets:
                if sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    sheet_rows = []
                    for row in sheet.iter_rows(values_only=True):
                        if any(row):
                            sheet_rows.append(list(row))
                    master_data[path].append({"sheet": sheet_name, "rows": sheet_rows})
                    
        elif file_type in ["docx", "doc"]:
            doc = Document(path)
            content = []
            for p in doc.paragraphs:
                if p.text.strip():
                    content.append({"type": "paragraph", "text": p.text})
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                content.append({"type": "table", "data": table_data})
            master_data[path] = content

    with open(output_json_path, "w", encoding="utf-8") as json_file:
        json.dump(master_data, json_file, indent=4, ensure_ascii=False)
        
    return f"Successfully generated parsed output and saved directly to {output_json_path}"


@tool
def execute_system_command(command: str) -> str:
    """
    Executes native shell toolchains (NPM format conversions or PeakRDL hardware compilations) 
    and surfaces back detailed STDOUT/STDERR logs for evaluation inside the Agent Loop.
    """
    try:
        result = subprocess.run(
            command,
            shell=True,
            text=True,
            capture_output=True,
            timeout=180
        )
        response = f"Exit Status Code: {result.returncode}\n"
        if result.stdout:
            response += f"--- STDOUT ---\n{result.stdout}\n"
        if result.stderr:
            response += f"--- STDERR ---\n{result.stderr}\n"
        return response
    except Exception as e:
        return f"Shell Execution faulted critical error structural trace: {str(e)}"

# ----------------------------------------------------
# 2. Deep Agent Factory Configuration (vLLM / Qwen 3.6)
# ----------------------------------------------------

def build_hardware_compiler_agent():
    # Hooking up to local high-throughput vLLM engine instance running Qwen 3.6
    model = ChatOpenAI(
        base_url="http://localhost:8000/v1",              # Update to your hosting vLLM URI endpoint
        api_key="EMPTY",                                  # Default token bypass for internal instances
        model="Qwen/Qwen3.6-35B-Instruct",                # Explicitly point to your active model tag
        temperature=0.0,                                  # Eliminate creative variance for syntax stability
        max_tokens=8192,                                  # Increased allocation space for large RDL files
        model_kwargs={
            "extra_body": {
                "top_p": 0.1                              # Pin precision on structural code elements
            }
        }
    )
    
    system_prompt = """You are an autonomous Hardware Register Engineering Copilot. Your target goal is translating raw documentation down into reliable C# peripheral models via the Antmicro PeakRDL-renode toolchain extension.

You have native environment tool capabilities via LangChain Deep Agents. Always organize files under the local `./workspace` directory structure.

### CRITICAL WORKFLOW STEPS:
1. **Plan Out Objectives**: Execute `write_todos` to create an initialization blueprint from the input variables.
2. **Document Extract**: Call `parse_raw_documents` using the user's explicit structural extraction rules (such as specified spreadsheet page blocks or sheet indexes).
3. **TOON Compression**: Execute the conversion package using the NPM runtime environment via `execute_system_command`:
   `npx jsonify-toonify ./workspace/master.json -o ./workspace/register.toon`
4. **RDL Compilation (Instruction-Guided)**: Read the conversion constraints provided in `./workspace/system_prompt_rdl.md`. Translate the contents of `register.toon` into a syntactically correct SystemRDL 2.0 configuration saved to `./workspace/register.rdl`.
5. **The PeakRDL-Renode Healing Loop**: 
   - Compile the C# models by running the official sub-exporter syntax via `execute_system_command`:
     `peakrdl renode ./workspace/register.rdl -o ./workspace/cs_output`
   
   - **Perceive & Observe**: Evaluate the shell output. PeakRDL utilizes the `systemrdl-compiler` layout. Errors will pinpoint exact lines matching structural rule failures (e.g., overlapping addresses, boundary errors, or invalid component configurations).
   
   - **Self-Correct Execution**: If the exit code is NOT 0, analyze the compiler logs. Update `./workspace/register.rdl` to resolve the syntax or alignment issue. Document the encountered failure alongside your fix inside `./workspace/README.md`.
   
   - **Iterate**: Re-run the `peakrdl renode` compilation command. Continue this loop until the exporter exits with status code 0.

6. **Final Sign-Off**: Once compilation succeeds, present the location of the generated C# assets inside `./workspace/cs_output/` along with the complete correction log in `README.md` to the user."""

    return create_deep_agent(
        model=model,
        tools=[parse_raw_documents, execute_system_command],
        system_prompt=system_prompt,
        name="hardware_rdl_deep_agent"
    )

# ----------------------------------------------------
# 3. Pipeline Initiation Interface 
# ----------------------------------------------------

if __name__ == "__main__":
    os.makedirs("workspace", exist_ok=True)
    
    # Pre-stage user's custom System Prompt MD definition path for RDL translation strategy
    if not os.path.exists("workspace/system_prompt_rdl.md"):
        with open("workspace/system_prompt_rdl.md", "w") as f:
            f.write("# RDL Transformation Rules\nTranslate TOON register representations into strict SystemRDL 2.0 format blocks.")

    # Initialize Deep Agent Orchestration Framework
    agent = build_hardware_compiler_agent()
    
    # Configure your structural multi-file task parameters below
    user_runtime_payload = {
        "messages": [
            {
                "role": "user",
                "content": """Please run the translation pipeline on these uploaded hardware targets:
                1. 'docs/soc_manual.pdf' -> target page range between 134 and 450 containing register arrays.
                2. 'docs/hardware_sheets.xlsm' -> look explicitly into data layout sheet names: ['SPI_CONFIG', 'I2C_MAP'].
                
                Track any unexpected syntax fixes required by PeakRDL-renode dynamically in a workspace README.md."""
            }
        ],
        "file_configs": [
            {"path": "docs/soc_manual.pdf", "type": "pdf", "pages": [134, 450]},
            {"path": "docs/hardware_sheets.xlsm", "type": "xlsm", "sheets": ["SPI_CONFIG", "I2C_MAP"]}
        ]
    }
    
    print("🚀 Initializing Deep Agent Framework execution loop running on vLLM + Qwen 3.6...")
    pipeline_state = agent.invoke(user_runtime_payload)
    print("\n🏁 Pipeline Completed successfully. Final Agent feedback:")
    print(pipeline_state["messages"][-1].content)
