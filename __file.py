"""
naive_rag_to_doc.py
- Input: large transcript.txt (~10-30 MB)
- Output: output.docx and output.pdf
- Assumes local LLM available via llama-cpp-python OR an ollama HTTP endpoint (toggle)
"""

import os
import textwrap
from pathlib import Path

# embeddings & vectorstore
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss

# LLM: choose llama-cpp bindings OR ollama REST
USE_OLLAMA = False  # set True to call ollama local server instead of llama-cpp
OLLAMA_URL = "http://localhost:11434/api/generate"

# For llama-cpp-python (if using)
try:
    from llama_cpp import Llama
except Exception:
    Llama = None

# output
from docx import Document
from reportlab.pdfgen import canvas

# ---------- Parameters ----------
TXT_PATH = "transcript.txt"
CHUNK_SIZE = 1500      # characters per chunk (adjust up/down)
CHUNK_OVERLAP = 200
EMBED_MODEL_NAME = "all-MiniLM-L6-v2"  # light and fast; change if you prefer
TOP_K = 6
LLAMA_MODEL_PATH = "/path/to/your/ggml-model.bin"  # for llama-cpp
OUTPUT_DOCX = "output.docx"
OUTPUT_PDF = "output.pdf"
# --------------------------------

def read_text(path):
    return Path(path).read_text(encoding="utf-8", errors="ignore")

def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    chunks = []
    i = 0
    n = len(text)
    while i < n:
        end = min(i + size, n)
        chunk = text[i:end]
        chunks.append(chunk.strip())
        i += size - overlap
    return chunks

def embed_chunks(model, chunks):
    emb = model.encode(chunks, show_progress_bar=True, convert_to_numpy=True)
    return emb.astype("float32")

def build_faiss_index(embeddings):
    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)  # inner product (we'll normalize vectors)
    faiss.normalize_L2(embeddings)
    index.add(embeddings)
    return index

def retrieve(query, embed_model, index, chunks, k=TOP_K):
    q_emb = embed_model.encode([query], convert_to_numpy=True).astype("float32")
    faiss.normalize_L2(q_emb)
    D, I = index.search(q_emb, k)
    results = []
    for idx in I[0]:
        if idx < len(chunks):
            results.append(chunks[idx])
    return results

def call_llama_local(prompt, max_tokens=2048, temp=0.2):
    if Llama is None:
        raise RuntimeError("llama-cpp-python not installed or import failed.")
    llm = Llama(model_path=LLAMA_MODEL_PATH)
    resp = llm.create(
        prompt=prompt,
        max_tokens=max_tokens,
        temperature=temp,
    )
    return resp.get("choices", [{}])[0].get("text", "")

def call_ollama(prompt):
    import requests, json
    data = {
        "model": "llama3",   # change to the model you have locally
        "prompt": prompt,
        "max_tokens": 2048,
        "temperature": 0.2
    }
    r = requests.post(OLLAMA_URL, json=data, timeout=120)
    r.raise_for_status()
    return r.json().get("output") or r.text

def assemble_generation_prompt(title, retrieved_chunks, instructions):
    header = f"Title: {title}\n\nInstructions: {instructions}\n\nContext:\n"
    # Stitch top-k retrieved chunks with separators
    ctx = "\n\n---\n\n".join(retrieved_chunks)
    footer = "\n\nNow produce a well-structured, well-documented document in markdown style including headings, a short summary (3-4 sentences), key timestamps or topics (if available), and a 'References' section indicating which chunks supported which points. Keep the document professional and suitable for export to Word/PDF."
    return header + ctx + footer

def save_docx(text, path):
    doc = Document()
    # naive conversion: split by markdown-ish headings
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        else:
            # simple paragraph wrapping
            doc.add_paragraph(line)
    doc.save(path)

def save_pdf_from_text(text, path):
    c = canvas.Canvas(path)
    width, height = c._pagesize
    y = height - 40
    for paragraph in text.split("\n\n"):
        lines = textwrap.wrap(paragraph, 100)
        for line in lines:
            c.drawString(40, y, line)
            y -= 12
            if y < 60:
                c.showPage()
                y = height - 40
        y -= 6
    c.save()

def main():
    print("Loading text...")
    text = read_text(TXT_PATH)
    print(f"Document size: {len(text)/1024/1024:.2f} MB")

    print("Chunking text...")
    chunks = chunk_text(text)
    print(f"{len(chunks)} chunks created (chunk_size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP})")

    print("Loading embedding model...")
    embed_model = SentenceTransformer(EMBED_MODEL_NAME)

    print("Embedding chunks (may take a while)...")
    embeddings = embed_chunks(embed_model, chunks)

    print("Building FAISS index...")
    index = build_faiss_index(embeddings)

    # Simple strategy: one generation iteration — ask model to produce final doc using top-k retrieved chunks
    title = "Auto-generated Document from Transcript"
    user_instructions = ("Produce a well-structured, documented report suitable for a technical reader. "
                         "Include a short executive summary, sectioned content, bullet points for key takeaways, and a references/notes section.")

    # Create a query that asks for a full document generation: you can refine this (e.g., generate section-by-section)
    query = "Create a professional document summarizing the whole transcript and explaining the main points."

    print("Retrieving top chunks...")
    retrieved = retrieve(query, embed_model, index, chunks, k=TOP_K)

    print("Assembling prompt for LLM...")
    prompt = assemble_generation_prompt(title, retrieved, user_instructions)

    print("Calling LLM...")
    if USE_OLLAMA:
        generated = call_ollama(prompt)
    else:
        generated = call_llama_local(prompt)

    print("Saving outputs...")
    save_docx(generated, OUTPUT_DOCX)
    save_pdf_from_text(generated, OUTPUT_PDF)

    print("Done. Saved:", OUTPUT_DOCX, OUTPUT_PDF)

if __name__ == "__main__":
    main()